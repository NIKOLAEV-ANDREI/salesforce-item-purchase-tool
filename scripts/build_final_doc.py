from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# The source file is intentionally kept ASCII-safe for the shell used in this
# workspace. Restore the UTF-8 Russian literals before python-docx writes them.
_original_add_run = Paragraph.add_run


def _add_run_with_fixed_encoding(self, text=None, style=None):
    if text is not None:
        try:
            text = text.encode('cp1251').decode('utf-8')
        except UnicodeError:
            pass
    return _original_add_run(self, text, style)


Paragraph.add_run = _add_run_with_fixed_encoding

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'deliverables' / 'Item_Purchase_Tool_Test_Assignment.docx'


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style='List Bullet')
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run(text))


def add_step(doc, text):
    paragraph = doc.add_paragraph(style='List Number')
    paragraph.paragraph_format.space_after = Pt(5)
    set_font(paragraph.add_run(text))


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f'Heading {level}')
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=(46, 116, 181))


def add_note(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run('Важно: ' + text)
    set_font(run, bold=True, color=(122, 90, 0))


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run('Item Purchase Tool'), size=22, bold=True, color=(11, 37, 69))
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
set_font(subtitle.add_run('Итоговый документ по тестовому заданию Salesforce'), size=12, color=(85, 85, 85))

add_heading(doc, '1. Реализованный функционал')
for text in [
    'Созданы объекты Item, Purchase и Purchase Line, включая связи, поля цены, количества, клиента и итогов покупки.',
    'На Account добавлена кнопка Item Purchase Tool. Она открывает Lightning Web Component и передаёт ID выбранного клиента.',
    'Инструмент показывает данные Account, каталог товаров, количество товаров в фильтре, фильтры Family и Type, поиск по названию и описанию, а также окно подробностей товара.',
    'Реализована корзина: добавление товаров, ограничение по доступному остатку, удаление из корзины, toast-уведомления и отображение итоговой суммы.',
    'Checkout выполняется на Apex-сервере в одной транзакции: остатки повторно проверяются с FOR UPDATE, создаются Purchase и Purchase Line, затем остатки уменьшаются. После успеха пользователь открывает стандартную страницу Purchase.',
    'Flow Purchase Calculate Totals пересчитывает TotalItems и GrandTotal покупки из строк покупки.',
    'Поле User.IsManager управляет доступом к кнопке New Item. Менеджер создаёт товар в форме, а URL изображения запрашивается через Unsplash API по названию товара.',
    'При изменении AvailableQuantity товара с положительного значения на 0 запускается Flow уведомления. Получатели берутся из Hierarchy Custom Setting; им отправляются Bell notification и email по шаблону.',
    'Созданы Permission Set, Remote Site Setting, CSP Trusted Site, Custom Notification Type и Email Template, необходимые для работы решения.'
]:
    add_bullet(doc, text)

add_heading(doc, '2. Как проверить')
add_heading(doc, 'Основной сценарий покупки', level=2)
for text in [
    'Войти в Salesforce и открыть Item Purchase Tool из записи Account либо по прямому URL компонента.',
    'Проверить отображение Name, Account Number и Industry клиента.',
    'Изменить Family или Type, ввести текст в поиск, открыть Details у товара.',
    'Добавить доступный товар в корзину. Товар с остатком 0 не должен добавляться.',
    'Открыть Cart, выполнить Checkout. Должна появиться Purchase, её строки, уменьшенный остаток и рассчитанные TotalItems / GrandTotal.'
]:
    add_step(doc, text)

add_heading(doc, 'Создание товара менеджером', level=2)
for text in [
    'У пользователя должен быть включён флаг IsManager.',
    'В Item Purchase Tool нажать New item, заполнить обязательные поля и сохранить.',
    'Убедиться, что в Item создан URL изображения. Для этого в Unsplash Settings должен быть заполнен Access Key.'
]:
    add_step(doc, text)

add_heading(doc, 'Проверка уведомления об остатке 0', level=2)
for text in [
    'В Setup открыть Custom Settings → Inventory Notification Settings и указать в Recipient User IDs Salesforce User ID получателя. Несколько ID разделяются точкой с запятой.',
    'Установить у любого товара Available Quantity в 1, затем изменить на 0.',
    'Проверить Bell notification в Salesforce и email получателя.'
]:
    add_step(doc, text)

add_note(doc, 'Для демонстрации в организации создана запись Account «Item Purchase Tool Demo Account».')

add_heading(doc, '3. Тесты и переносимый пакет')
for text in [
    'Проверочное развёртывание метаданных завершилось успешно: 41 компонент, 9 Apex-тестов, без ошибок.',
    'В deliverables приложен ItemPurchaseTool-Unmanaged.zip: Metadata API ZIP с package.xml. Он успешно прошёл dry-run deploy в текущей организации.',
    'Создан пользователь: Email = dev@truesolv.com, Username = dev+item-purchase-tool@truesolv.com, профиль Standard User, назначен ItemPurchaseToolAccess.'
]:
    add_bullet(doc, text)

add_heading(doc, '4. Известные ограничения')
for text in [
    'Логин dev@truesolv.com оказался глобально занят в другой Salesforce-организации. Поэтому у созданного пользователя сохранён требуемый Email, но Username использует уникальный суффикс.',
    'Email-уведомление содержит общий текст. Название товара содержится в Bell notification: Salesforce не разрешает передавать кастомный WhatId в template email, когда получатель является User.',
    'Для Unsplash нужен действующий Access Key и действует лимит demo-приложения. Ключ хранится только в настройках Salesforce и не включён в исходный код или ZIP.',
    'Подготовлен валидированный Metadata API ZIP. Нативный first-generation unmanaged package из Package Manager не создан автоматически, поскольку этот объект доступен только через Salesforce Package Manager UI в данной организации.'
]:
    add_bullet(doc, text)

doc.add_paragraph()
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run('Salesforce Item Purchase Tool'), size=9, color=(85, 85, 85))

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
